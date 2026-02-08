from __future__ import annotations

import os
import re
import threading
from datetime import datetime
from typing import Dict, List, Tuple

import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config.app_config import load_config, get_config_path
from services.timeweb_ai import stream_chat
from services import webdav_sync


SYSTEM_PROMPT = (
    "Ты — эксперт по микробиологическому мониторингу в медицинских организациях. "
    "Твои выводы и рекомендации должны опираться на требования и подходы, изложенные в: "
    "СанПиН 3.3686 и МР 3.1.0346-24 «Организация и проведение микробиологического мониторинга в медицинских организациях». "
    "Работай только с данными, представленными в отчётах. Если чего-то не хватает — явно укажи."
)


def _docx_to_text(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts)


def _build_user_prompt(files: List[str], texts: List[str]) -> str:
    if len(files) == 1:
        header = "Сделай анализ одного микробиологического отчёта."
    else:
        header = "Сделай сравнение микробиологических отчётов (динамика, изменения, риски)."

    body_parts = []
    for path, text in zip(files, texts):
        name = os.path.basename(path)
        body_parts.append(f"=== ОТЧЁТ: {name} ===\n{text}")

    return header + "\n\n" + "\n\n".join(body_parts)


def build_ai_analysis_screen(main_frame, build_header, go_back_callback, archive_dir: str, ai_save_dir: str):
    for w in main_frame.winfo_children():
        w.destroy()

    build_header(main_frame, back_callback=go_back_callback)
    webdav_sync.sync_down(webdav_sync.get_default_local_root())

    tk.Label(
        main_frame,
        text="🤖 AI анализ",
        font=("Segoe UI", 18, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(18, 8))

    root = tk.Frame(main_frame, bg="#f4f6f8")
    root.pack(expand=True, fill="both", padx=18, pady=10)

    left = tk.Frame(root, bg="#f4f6f8", width=360)
    left.pack(side="left", fill="y")
    left.pack_propagate(False)

    right = tk.Frame(root, bg="#f4f6f8")
    right.pack(side="right", expand=True, fill="both")

    tk.Label(left, text="Выберите отчёты", font=("Segoe UI", 11, "bold"), bg="#f4f6f8").pack(anchor="w")
    listbox = tk.Listbox(left, selectmode="extended", width=42, height=22)
    listbox.pack(fill="both", expand=True, pady=(6, 8))

    btns = tk.Frame(left, bg="#f4f6f8")
    btns.pack(fill="x")

    btn_refresh = ttk.Button(btns, text="Обновить")
    btn_refresh.pack(side="left", expand=True, fill="x", padx=(0, 6))

    btn_analyze = ttk.Button(btns, text="AI анализ")
    btn_analyze.pack(side="left", expand=True, fill="x")

    status = tk.Label(left, text="", bg="#f4f6f8", fg="#444", anchor="w")
    status.pack(fill="x", pady=(6, 0))

    # right side: output
    text_box = tk.Text(right, font=("Times New Roman", 14), wrap="word")
    text_box.pack(expand=True, fill="both")

    btn_save = ttk.Button(right, text="Сохранить в DOCX")
    btn_save.pack(pady=(8, 0))
    btn_save.config(state="disabled")

    file_map: Dict[int, str] = {}
    last_text: List[str] = [""]
    stop_stream = threading.Event()
    text_box.tag_configure("h1", font=("Times New Roman", 16, "bold"))
    text_box.tag_configure("body", font=("Times New Roman", 14))
    text_box.tag_configure("latin", font=("Times New Roman", 14, "italic"))

    def _clean_markdown(text: str) -> str:
        # Lightweight cleanup for markdown markers:
        # - remove '---' separator lines later
        # - strip markdown emphasis (*, **), but keep text
        t = text.replace("\r\n", "\n")
        t = t.replace("**", "").replace("__", "").replace("`", "")
        t = t.replace("### ", "").replace("## ", "").replace("# ", "")
        t = t.replace("* ", "- ").replace("- ", "- ")
        t = t.replace("*", "")
        return t

    _re_sep = re.compile(r"^\\s*[-—_]{3,}\\s*$")
    _re_heading_num = re.compile(r"^\\s*\\d+(?:\\.\\d+)*[.)]\\s+\\S+")
    _re_latin_binomial = re.compile(r"\b[A-Z][a-z]{2,}\s+(?:[a-z]{2,}|spp\.|sp\.)\b")
    _re_latin_abbrev = re.compile(r"\b[A-Z]\.\s*[a-z]{2,}\b")

    def _is_separator_line(line: str) -> bool:
        return bool(_re_sep.match(line or ""))

    def _latin_spans(s: str) -> List[Tuple[int, int]]:
        spans: List[Tuple[int, int]] = []
        for m in _re_latin_binomial.finditer(s):
            spans.append((m.start(), m.end()))
        for m in _re_latin_abbrev.finditer(s):
            spans.append((m.start(), m.end()))
        if not spans:
            return spans
        spans.sort()
        merged = [spans[0]]
        for a, b in spans[1:]:
            la, lb = merged[-1]
            if a <= lb:
                merged[-1] = (la, max(lb, b))
            else:
                merged.append((a, b))
        return merged

    def _insert_with_latin_tags(line: str) -> None:
        if stop_stream.is_set() or not text_box.winfo_exists():
            return

        spans = _latin_spans(line)
        if not spans:
            text_box.insert(tk.END, line + "\n", "body")
            return

        pos = 0
        for a, b in spans:
            if a > pos:
                text_box.insert(tk.END, line[pos:a], "body")
            text_box.insert(tk.END, line[a:b], ("body", "latin"))
            pos = b
        if pos < len(line):
            text_box.insert(tk.END, line[pos:], "body")
        text_box.insert(tk.END, "\n", "body")

    def _is_heading(line: str, *, in_list_context: bool) -> bool:
        s = (line or "").strip()
        if not s:
            return False

        if s.lower() in ("выводы:", "рекомендации:", "имеется:", "отсутствует:"):
            return False

        if s.startswith("Анализ") or s.startswith("Сравнение") or s.startswith("Краткое резюме"):
            return True

        if (not in_list_context) and _re_heading_num.match(s):
            # headings are usually short; long lines are likely list items
            words = len(s.split())
            if words <= 12 and len(s) <= 90:
                return True

        return False

    def set_status(msg: str):
        status.config(text=msg)

    def load_files():
        listbox.delete(0, tk.END)
        file_map.clear()
        webdav_sync.sync_down(webdav_sync.get_default_local_root())

        if not os.path.exists(archive_dir):
            set_status(f"Архив не найден: {archive_dir}")
            return

        items = []
        for root_dir, _, files in os.walk(archive_dir):
            for f in files:
                if f.lower().endswith(".docx"):
                    full = os.path.join(root_dir, f)
                    rel = os.path.relpath(full, archive_dir)
                    items.append((rel, full))

        items.sort(key=lambda x: x[0].lower())
        for idx, (rel, full) in enumerate(items):
            listbox.insert(tk.END, rel)
            file_map[idx] = full

        set_status(f"Файлов: {len(items)}")

    def _append_text(chunk: str):
        if stop_stream.is_set() or not text_box.winfo_exists():
            return
        text_box.insert(tk.END, _clean_markdown(chunk), "body")
        text_box.see(tk.END)

    def analyze_selected():
        cfg = load_config()
        api_key = cfg.get("ai_api_key", "").strip()
        base_url = cfg.get("ai_base_url", "").strip()
        agent_id = cfg.get("ai_agent_id", "").strip()

        if not (api_key and base_url and agent_id):
            messagebox.showerror(
                "Настройки AI",
                "Заполните ai_api_key, ai_base_url и ai_agent_id в config/app_config.json.\n\n"
                f"Текущий путь конфига:\n{get_config_path()}"
            )
            return

        sel = listbox.curselection()
        if not sel:
            messagebox.showinfo("AI анализ", "Выберите один или несколько отчётов.")
            return

        paths = [file_map[i] for i in sel if i in file_map]
        if not paths:
            return

        texts = []
        try:
            for p in paths:
                texts.append(_docx_to_text(p))
        except Exception as e:
            messagebox.showerror("AI анализ", f"Не удалось прочитать DOCX:\n{e}")
            return

        user_prompt = _build_user_prompt(paths, texts)
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]

        if text_box.winfo_exists():
            text_box.delete(1.0, tk.END)
        btn_save.config(state="disabled")
        last_text[0] = ""
        set_status("Анализируем…")

        def worker():
            try:
                for chunk in stream_chat(
                    base_url=base_url,
                    api_key=api_key,
                    agent_id=agent_id,
                    messages=messages,
                ):
                    if stop_stream.is_set() or not text_box.winfo_exists():
                        break
                    main_frame.after(0, _append_text, chunk)
                    last_text[0] += chunk
            except Exception as e:
                main_frame.after(0, messagebox.showerror, "AI анализ", str(e))
            finally:
                def finish():
                    if stop_stream.is_set() or not text_box.winfo_exists():
                        return
                    if last_text[0].strip():
                        btn_save.config(state="normal")
                        # re-render with heading styles after streaming
                        text_box.delete(1.0, tk.END)
                        cleaned = _clean_markdown(last_text[0])
                        in_list_context = False
                        for raw in cleaned.splitlines():
                            line = raw.rstrip()
                            if _is_separator_line(line):
                                continue  # no '---' lines
                            s = line.strip()
                            if not s:
                                text_box.insert(tk.END, "\n", "body")
                                continue

                            if s.lower() in ("выводы:", "рекомендации:"):
                                in_list_context = True
                                text_box.insert(tk.END, s + "\n", "body")
                                continue

                            if _is_heading(s, in_list_context=in_list_context):
                                in_list_context = False
                                text_box.insert(tk.END, s + "\n", "h1")
                                continue

                            _insert_with_latin_tags(s)
                    set_status("Готово")
                main_frame.after(0, finish)

        threading.Thread(target=worker, daemon=True).start()

    def _on_text_destroy(_):
        stop_stream.set()

    text_box.bind("<Destroy>", _on_text_destroy)

    def save_docx():
        if not last_text[0].strip():
            return

        os.makedirs(ai_save_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"AI_анализ_{ts}.docx"
        out_path = os.path.join(ai_save_dir, filename)

        doc = Document()
        title = doc.add_heading("AI анализ микробиологических отчётов", level=1)
        for run in title.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)

        cleaned = _clean_markdown(last_text[0])
        in_list_context = False
        for raw in cleaned.splitlines():
            line = raw.rstrip()
            if _is_separator_line(line):
                continue  # no '---' lines
            s = line.strip()
            if not s:
                doc.add_paragraph("")
                continue

            if s.lower() in ("выводы:", "рекомендации:"):
                in_list_context = True
                p = doc.add_paragraph(s)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                continue

            if _is_heading(s, in_list_context=in_list_context):
                in_list_context = False
                p = doc.add_paragraph()
                run = p.add_run(s)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(16)
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            spans = _latin_spans(s)
            if not spans:
                run = p.add_run(s)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                continue

            pos = 0
            for a, b in spans:
                if a > pos:
                    run = p.add_run(s[pos:a])
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                run = p.add_run(s[a:b])
                run.italic = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                pos = b
            if pos < len(s):
                run = p.add_run(s[pos:])
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
        try:
            doc.save(out_path)
            webdav_sync.upload_file(out_path, webdav_sync.get_default_local_root())
            messagebox.showinfo("AI анализ", f"Сохранено:\n{out_path}")
        except Exception as e:
            messagebox.showerror("AI анализ", f"Не удалось сохранить DOCX:\n{e}")

    btn_refresh.config(command=load_files)
    btn_analyze.config(command=analyze_selected)
    btn_save.config(command=save_docx)

    load_files()
