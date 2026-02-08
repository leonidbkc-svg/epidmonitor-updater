import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Inches
import tempfile
import os
import shutil
import calendar
import datetime
from config.app_config import load_config, save_config, get_config_path
import getpass
from services.tg_counters import read_latest_numbers
from services import webdav_sync
from screens.tg_exam_stats import open_tg_exam_stats
from screens.testing_menu import open_testing_menu


from screens.documents import build_documents_screen
from screens.vaccination_stats import open_vaccination_stats
from screens.testing import open_testing_screen


from screens.swab_monitoring import build_swab_monitoring_screen
from PIL import Image, ImageTk

from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

from analysis.loci import LOCUS_COLORS
from screens.photo_rounds import build_photo_rounds_screen
from screens.ai_analysis import build_ai_analysis_screen
from screens.microbes_tab import build_microbes_tab
from screens.loci_tab import build_loci_tab
from screens.resistance_tab import build_resistance_tab
from utils.charts import (
    bar_with_value_labels,
    barh_with_value_labels,
    reset_canvas,
    stacked_barh,
)


APP_PASSWORD = "2"  # ← поменяешь на свой

# ======================================================
# ВАКЦИНАЦИЯ — СТАТИСТИКА
# ======================================================

# База сервера вакцинации (VPS + домен + https)
VACC_SERVER_BASE = "https://vaccine.epid-test.ru"

# Пароль входа во вкладку в твоём приложении (локальная защита UI)
VACC_TAB_PASSWORD = "1234"   # можешь оставить как есть или поменять

# PIN для API на сервере (тот же, что ADMIN_PIN в systemd)
VACC_ADMIN_PIN = "1234"

TEST_SERVER_BASE = "http://147.45.163.136"
TEST_ADMIN_PIN = ""





# ======================================================
# ОБЩИЙ АРХИВ (СЕТЬ) + ФОЛБЭК НА ЛОКАЛЬНЫЙ
# ======================================================
# На скрине у тебя именно админ-ресурс c$:
# \\192.168.137.17\c$\EpidArchive
#
# В ИДЕАЛЕ (для всех пользователей) сделать обычную шару:
# \\192.168.137.17\EpidArchive
#
# Пока ставим как у тебя на скрине:
NETWORK_ROOT = r"\\192.168.137.17\c$\EpidArchive"
WEBDAV_ERROR_SHOWN = False

def pick_data_root() -> str:
    """
    1) Если доступен сетевой архив — используем его
    2) Если нет — падаем назад в AppData пользователя
    """
    try:
        cfg = load_config()
        webdav_url = (cfg.get("webdav_url") or "").strip()
        if webdav_url:
            raise RuntimeError("force_local_for_webdav")

        if os.path.exists(NETWORK_ROOT):
            os.makedirs(NETWORK_ROOT, exist_ok=True)

            # проверка записи (чтобы не было сюрпризов)
            test_path = os.path.join(NETWORK_ROOT, "_write_test.tmp")
            with open(test_path, "w", encoding="utf-8") as f:
                f.write("ok")
            os.remove(test_path)

            return NETWORK_ROOT
    except Exception:
        pass

    # fallback локально
    local_root = os.path.join(
        os.environ.get("APPDATA", os.path.expanduser("~")),
        "EpidMonitor"
    )
    os.makedirs(local_root, exist_ok=True)
    return local_root

DATA_ROOT = pick_data_root()

# ВАЖНО: эти папки должны совпадать для сохранения и чтения архива
ARCHIVE_DIR = os.path.join(DATA_ROOT, "reports_archive")
SWABS_ARCHIVE_DIR = r"\\w00164\e$\ДОКУМЕНТЫ ОТДЕЛА\2026\ИСМП 2026\Протоколы исследований\Автоматические отчеты\Word"
PHOTO_ROUNDS_DIR = os.path.join(DATA_ROOT, "archive", "photo_rounds")
DOCUMENTS_DIR = os.path.join(DATA_ROOT, "documents")
AI_REPORTS_DIR = os.path.join(ARCHIVE_DIR, "AI")

os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(PHOTO_ROUNDS_DIR, exist_ok=True)
os.makedirs(ARCHIVE_DIR, exist_ok=True)
try:
    os.makedirs(SWABS_ARCHIVE_DIR, exist_ok=True)
except Exception:
    # сеть/права могут быть недоступны на старте
    pass






report_state = {
    "microbes": None,
    "loci": None,
    "resistance": None
}


def normalize_gram(name: str) -> str:
    name = (name or "").lower()
    if "полож" in name:
        return "Gram+"
    if "отриц" in name:
        return "Gram-"
    if "гриб" in name or "fung" in name:
        return "Грибы"
    return "Не классифицировано"


def _get_windows_user() -> str:
    try:
        name = getpass.getuser()
        if name:
            return name
    except Exception:
        pass
    return os.environ.get("USERNAME") or "пользователь"


# ======================================================
# СОХРАНЕНИЕ ОТЧЁТА (ОБЁРТКА)
# ======================================================

def save_report():
    missing = []

    if report_state["microbes"] is None:
        missing.append("Микроорганизмы")
    if report_state["loci"] is None:
        missing.append("Локусы")
    if report_state["resistance"] is None:
        missing.append("Резистентность")

    if missing:
        text = (
            "Не заполнены вкладки:\n\n"
            + "\n".join(f"• {m}" for m in missing)
            + "\n\nТочно сохранить отчёт?"
        )
        if not messagebox.askyesno("Неполный отчёт", text):
            return

    do_save_report_docx()






def do_save_report_docx():
    # ---------- ИМЯ ФАЙЛА ----------
    base_name = f"{current_department}_{current_month}_{current_year}".replace(" ", "_")

    file = filedialog.asksaveasfilename(
        defaultextension=".docx",
        initialfile=base_name + ".docx",
        filetypes=[("Word документ", "*.docx")]
    )
    if not file:
        return

    dep_name = (current_department or "Без_отделения").replace("/", "_").replace("\\", "_")
    dep_dir = os.path.join(ARCHIVE_DIR, dep_name)
    os.makedirs(dep_dir, exist_ok=True)

    archive_path = os.path.join(dep_dir, base_name + ".docx")

    # ---------- DOCX ----------
    doc = Document()
    tmp_dir = tempfile.mkdtemp()

    # ==================================================
    # ЗАГОЛОВОК
    # ==================================================
    doc.add_heading("Эпидемиологический отчёт", level=1)

    p = doc.add_paragraph()
    p.add_run("Отделение: ").bold = True
    p.add_run(str(current_department or "") + "\n")
    p.add_run("Период: ").bold = True
    p.add_run(f"{current_month or ''} {current_year or ''}")

    # ==================================================
    # 1. МИКРООРГАНИЗМЫ + GRAM
    # ==================================================
    if report_state.get("microbes"):
        r = report_state["microbes"]

        doc.add_heading("1. Микроорганизмы", level=2)

        total = int(r.get("total", 0))
        doc.add_paragraph(f"Всего выделено микроорганизмов: {total}.")

        microbes = sorted(
            r.get("microbes", []),
            key=lambda x: int(x.get("count", 0)),
            reverse=True
        )

        for m in microbes:
            doc.add_paragraph(
                f"{m.get('microbe','')} — {m.get('count',0)} ({m.get('percent',0)}%)",
                style="List Bullet"
            )

        # ---------- GRAM ----------
        doc.add_paragraph("Распределение по Gram")

        gram_map = {
            "Gram+": 0,
            "Gram-": 0,
            "Грибы": 0,
            "Не классифицировано": 0
        }

        for m in microbes:
            gram = m.get("gram", "Не классифицировано")
            if gram not in gram_map:
                gram = "Не классифицировано"
            gram_map[gram] += int(m.get("count", 0))

        for gram, count in gram_map.items():
            if count and total:
                percent = round(count / total * 100, 1)
                doc.add_paragraph(
                    f"{gram}: {count} ({percent}%)",
                    style="List Bullet"
                )

        # ---------- ГРАФИК: микроорганизмы ----------
        img_microbes = os.path.join(tmp_dir, "microbes_counts.png")
        fig = Figure(figsize=(8.5, 5), dpi=130)
        ax = fig.add_subplot(111)

        labels = [m["microbe"] for m in microbes][::-1]
        values = [int(m["count"]) for m in microbes][::-1]

        barh_with_value_labels(
            ax,
            labels,
            values,
            title="Микроорганизмы (количество)",
            xlabel="Количество"
        )

        fig.tight_layout()
        fig.savefig(img_microbes)
        doc.add_picture(img_microbes, width=Inches(6.5))

        # ---------- ГРАФИК: Gram ----------
        img_gram = os.path.join(tmp_dir, "gram.png")
        fig2 = Figure(figsize=(7.5, 4.5), dpi=130)
        ax2 = fig2.add_subplot(111)

        g_labels = [k for k, v in gram_map.items() if v > 0]
        g_values = [gram_map[k] for k in g_labels]

        bar_with_value_labels(
            ax2,
            g_labels,
            g_values,
            title="Распределение по Gram"
        )

        fig2.tight_layout()
        fig2.savefig(img_gram)
        doc.add_picture(img_gram, width=Inches(6.5))

    # ==================================================
    # 2. ЛОКУСЫ
    # ==================================================
    if report_state.get("loci"):
        r = report_state["loci"]
        groups = r.get("groups", [])

        doc.add_heading("2. Локусы", level=2)

        for g in groups:
            doc.add_paragraph(
                f"{g.get('group','')} — {g.get('count',0)} ({g.get('percent',0)}%)",
                style="List Bullet"
            )
            for it in g.get("items", []) or []:
                doc.add_paragraph(
                    f"{it.get('name','')} — {it.get('count',0)} ({it.get('percent',0)}%)",
                    style="List Bullet 2"
                )

        has_items = any(g.get("items") for g in groups)
        if has_items:
            img_loci = os.path.join(tmp_dir, "loci_stacked.png")

            groups_sorted = sorted(groups, key=lambda x: x["count"], reverse=True)
            group_names = [g["group"] for g in groups_sorted]

            totals = {}
            for g in groups_sorted:
                for it in g.get("items", []):
                    totals[it["name"]] = totals.get(it["name"], 0) + int(it["count"])

            loci_order = sorted(totals.keys(), key=lambda k: totals[k], reverse=True)

            fig = Figure(figsize=(10, 5.2), dpi=130)
            ax = fig.add_subplot(111)

            group_totals = [g["count"] for g in groups_sorted]
            max_total = max(group_totals) if group_totals else 0
            groups_items = [g.get("items", []) for g in groups_sorted]
            stacked_barh(
                ax,
                group_names,
                loci_order,
                groups_items,
                colors_map=LOCUS_COLORS
            )

            for i, total in enumerate(group_totals):
                ax.text(
                    total + (max_total * 0.01 + 0.5 if max_total else 0.5),
                    i,
                    str(total),
                    va="center",
                    fontsize=9,
                    fontweight="bold"
                )

            ax.set_title("Структура биоматериалов по локусам")
            ax.set_xlabel("Количество")
            ax.set_xlim(0, max_total * 1.2 if max_total else 10)

            ax.legend(
                title="Локус",
                fontsize=8,
                bbox_to_anchor=(1.02, 1),
                loc="upper left"
            )

            fig.tight_layout()
            fig.savefig(img_loci)
            doc.add_picture(img_loci, width=Inches(6.5))

    # ==================================================
    # 3. РЕЗИСТЕНТНОСТЬ
    # ==================================================
    if report_state.get("resistance"):
        r = report_state["resistance"]

        doc.add_heading("3. Резистентность", level=2)

        microbes = sorted(r.get("microbes", []), key=lambda x: x["r_percent"], reverse=True)
        antibiotics = sorted(r.get("antibiotics", []), key=lambda x: x["r_percent"], reverse=True)

        doc.add_paragraph("Резистентность по микроорганизмам")
        for m in microbes:
            doc.add_paragraph(
                f'{m["microbe"]} — R {m["r_percent"]}% ({m["r_count"]}/{m["total"]})',
                style="List Bullet"
            )

        doc.add_paragraph("Резистентность по антибиотикам")
        for a in antibiotics:
            doc.add_paragraph(
                f'{a["antibiotic"]} — R {a["r_percent"]}% ({a["r_count"]}/{a["total"]})',
                style="List Bullet"
            )

        # --- график микроорганизмы ---
        img_m = os.path.join(tmp_dir, "res_microbes.png")
        fig_m = Figure(figsize=(8.5, 5), dpi=130)
        ax_m = fig_m.add_subplot(111)

        labels = [m["microbe"] for m in microbes][::-1]
        values = [m["r_percent"] for m in microbes][::-1]

        barh_with_value_labels(
            ax_m,
            labels,
            values,
            title="Резистентность по микроорганизмам",
            xlabel="R (%)",
            value_fmt=lambda v: str(int(v)) if float(v).is_integer() else str(v)
        )

        fig_m.tight_layout()
        fig_m.savefig(img_m)
        doc.add_picture(img_m, width=Inches(6.5))

        # --- график антибиотики ---
        img_a = os.path.join(tmp_dir, "res_antibiotics.png")
        fig_a = Figure(figsize=(8.5, 5), dpi=130)
        ax_a = fig_a.add_subplot(111)

        labels = [a["antibiotic"] for a in antibiotics][::-1]
        values = [a["r_percent"] for a in antibiotics][::-1]

        barh_with_value_labels(
            ax_a,
            labels,
            values,
            title="Резистентность по антибиотикам",
            xlabel="R (%)",
            value_fmt=lambda v: str(int(v)) if float(v).is_integer() else str(v)
        )

        fig_a.tight_layout()
        fig_a.savefig(img_a)
        doc.add_picture(img_a, width=Inches(6.5))

    # ==================================================
    # СОХРАНЕНИЕ
    # ==================================================
    try:
        doc.save(file)
        doc.save(archive_path)

        # синхронизируем сохранённый отчёт в WebDAV (если настроен)
        webdav_sync.upload_file(archive_path, DATA_ROOT)

        messagebox.showinfo(
            "Готово",
            "DOCX-отчёт успешно сохранён\n\n"
            "• в выбранную папку\n"
            "• в архив программы"
        )
    except Exception as e:
        messagebox.showerror("Ошибка сохранения", str(e))



    


# ======================================================
# ROOT (ОДИН РАЗ) + ЗАСТАВКА
# ======================================================

root = tk.Tk()
root.withdraw()   # скрываем главное окно
root.title("ЭпидМонитор")

splash_window = None  # будет создан в show_splash_screen


def show_splash_screen():
    """Показ заставки перед запуском программы (работает в EXE)."""
    global splash_window
    splash_window = tk.Toplevel(root)
    splash_window.overrideredirect(True)
    splash_w = 420
    splash_h = 300
    splash_window.update_idletasks()
    screen_w = splash_window.winfo_screenwidth()
    screen_h = splash_window.winfo_screenheight()
    x = max(0, int((screen_w - splash_w) / 2))
    y = max(0, int((screen_h - splash_h) / 2))
    splash_window.geometry(f"{splash_w}x{splash_h}+{x}+{y}")

    frame = tk.Frame(splash_window, bg="white")
    frame.pack(expand=True, fill="both")

    # Логотип
    try:
        img = Image.open("logo.png").resize((150, 150), Image.LANCZOS)
        logo_img = ImageTk.PhotoImage(img)
        lbl = tk.Label(frame, image=logo_img, bg="white")
        lbl.image = logo_img
        lbl.pack(pady=(30, 10))
    except Exception:
        pass

    # Надписи
    tk.Label(
        frame,
        text="ЭпидМонитор",
        font=("Segoe UI", 16, "bold"),
        bg="white"
    ).pack()

    tk.Label(
        frame,
        text="Загрузка системы...",
        font=("Segoe UI", 11),
        bg="white",
        fg="#555"
    ).pack(pady=(10, 20))


# ======================================================
# СТИЛИ
# ======================================================

style = ttk.Style()
style.theme_use("default")
style.configure("Main.TButton", font=("Segoe UI", 11), padding=10)
style.configure("Secondary.TButton", font=("Segoe UI", 10), padding=8)
style.configure("TNotebook.Tab", font=("Segoe UI", 10), padding=(12, 6))


# ======================================================
# ОСНОВНОЙ КОНТЕЙНЕР
# ======================================================

main_frame = tk.Frame(root, bg="#f4f6f8")
main_frame.pack(expand=True, fill="both")


# ======================================================
# КОНТЕКСТ ОТЧЁТА
# ======================================================

current_department = None
current_month = None
current_year = None


# ======================================================
# ЛОГИН
# ======================================================

def build_login_screen():
    for w in main_frame.winfo_children():
        w.destroy()

    tk.Label(
        main_frame,
        text="ЭпидМонитор",
        font=("Segoe UI", 22, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(80, 10))

    try:
        img = Image.open("logo.png").resize((120, 120), Image.LANCZOS)
        logo_img = ImageTk.PhotoImage(img)
        lbl = tk.Label(main_frame, image=logo_img, bg="#f4f6f8")
        lbl.image = logo_img
        lbl.pack(pady=(0, 30))
    except Exception:
        pass

    tk.Label(
        main_frame,
        text="Введите пароль для входа",
        font=("Segoe UI", 12),
        bg="#f4f6f8"
    ).pack(pady=(0, 8))

    pwd_var = tk.StringVar()

    pwd_entry = ttk.Entry(
        main_frame,
        textvariable=pwd_var,
        show="*",
        width=30,
        font=("Segoe UI", 11)
    )
    pwd_entry.pack(pady=(0, 15))
    pwd_entry.focus()

    def check_password(event=None):
        if pwd_var.get() == APP_PASSWORD:
            user = _get_windows_user()
            messagebox.showinfo("Добро пожаловать", f"Добро пожаловать, {user}!")
            root.unbind("<Return>")
            build_start_screen()
        else:
            messagebox.showerror("Ошибка", "Неверный пароль")
            pwd_var.set("")
            pwd_entry.focus()

    ttk.Button(
        main_frame,
        text="Войти",
        style="Main.TButton",
        command=check_password
    ).pack(padx=200, fill="x")

    root.bind("<Return>", check_password)

# ======================================================
# СТАРТОВЫЙ ЭКРАН
# ======================================================

def build_start_screen():
    # очистка экрана
    for w in main_frame.winfo_children():
        w.destroy()

    # верхняя панель со "Настройками"
    top_bar = tk.Frame(main_frame, bg="#f4f6f8")
    top_bar.pack(fill="x", pady=(10, 0), padx=10)

    ttk.Button(
        top_bar,
        text="⚙ Настройки",
        style="Secondary.TButton",
        command=lambda: open_settings_dialog(main_frame)
    ).pack(side="right")

    global WEBDAV_ERROR_SHOWN
    cfg = load_config()
    if (cfg.get("webdav_url") or "").strip():
        ok = webdav_sync.sync_down(DATA_ROOT)
        if not ok and not WEBDAV_ERROR_SHOWN:
            WEBDAV_ERROR_SHOWN = True
            err = webdav_sync.get_last_error() or "Не удалось синхронизировать WebDAV."
            msg = (
                "Не удалось подключиться к WebDAV-хранилищу.\n"
                "Программа использует локальный архив:\n"
                f"{DATA_ROOT}\n\n"
                "Ошибка:\n"
                f"{err}"
            )
            messagebox.showwarning("WebDAV", msg)

    # ---------- ЗАГОЛОВОК ----------
    tk.Label(
        main_frame,
        text="ЭпидМонитор",
        font=("Segoe UI", 22, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(60, 8))

    # ---------- ЛОГО ----------
    try:
        img = Image.open("logo.png").resize((140, 140), Image.LANCZOS)
        logo_img = ImageTk.PhotoImage(img)

        lbl = tk.Label(main_frame, image=logo_img, bg="#f4f6f8")
        lbl.image = logo_img
        lbl.pack(pady=(10, 20))
    except Exception:
        pass

    # ---------- ОПИСАНИЕ ----------
    tk.Label(
        main_frame,
        text="Система эпидемиологического мониторинга",
        font=("Segoe UI", 12),
        bg="#f4f6f8",
        fg="#4b5563"
    ).pack(pady=(0, 40))

    # ---------- КНОПКИ ----------
    buttons = tk.Frame(main_frame, bg="#f4f6f8")
    buttons.pack()

    ttk.Button(
        buttons,
        text="📝 Создать отчёт",
        style="Main.TButton",
        command=show_create_report_screen
    ).pack(fill="x", pady=6)

    ttk.Button(
        buttons,
        text="🧪 Мониторинг смывов",
        style="Main.TButton",
        command=lambda: build_swab_monitoring_screen(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen
        )
    ).pack(fill="x", pady=6)

    ttk.Button(
        buttons,
        text="📷 Фотоотчёт обходов",
        style="Main.TButton",
        command=lambda: build_photo_rounds_screen(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen
        )
    ).pack(fill="x", pady=6)

    ttk.Button(
        buttons,
        text="📁 Архив отчётов",
        style="Main.TButton",
        command=build_archive_choice_screen
    ).pack(fill="x", pady=6)

    ttk.Button(
        buttons,
        text="📄 Документы",
        style="Main.TButton",
        command=lambda: build_documents_screen(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen
        )
    ).pack(fill="x", pady=6)

    # ---------- СТАТИСТИКА ВАКЦИНАЦИИ ----------
    ttk.Button(
        buttons,
        text="💉 Статистика вакцинации",
        style="Main.TButton",
        command=lambda: open_vaccination_stats(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen,
            server_base_url=VACC_SERVER_BASE,
            access_password=VACC_TAB_PASSWORD,
            admin_pin=VACC_ADMIN_PIN,
            refresh_ms=30000,
        )
    ).pack(fill="x", pady=6)

    # ---------- ТЕСТИРОВАНИЕ (МЕНЮ) ----------
    def _open_tg_exam_menu():
        cfg = load_config()
        tg_exam_base = (cfg.get("tg_exam_base_url", "") or "").strip()
        tg_exam_key = (cfg.get("tg_exam_report_api_key", "") or "").strip()
        if not tg_exam_base:
            tg_exam_base = "http://147.45.163.136"
        open_testing_menu(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen,
            data_root=DATA_ROOT,
            base_url=tg_exam_base,
            report_api_key=tg_exam_key,
            test_server_base=tg_exam_base or TEST_SERVER_BASE,
            test_admin_pin=TEST_ADMIN_PIN,
            vacc_server_base=VACC_SERVER_BASE,
            vacc_admin_pin=VACC_ADMIN_PIN,
            qr_image_path="qr_epid_test.png",
        )

    ttk.Button(
        buttons,
        text="📋 Тестирование",
        style="Main.TButton",
        command=lambda: _open_tg_exam_menu()
    ).pack(fill="x", pady=6)

    # ---------- AI АНАЛИЗ ----------
    ttk.Button(
        buttons,
        text="🤖 AI анализ",
        style="Main.TButton",
        command=lambda: build_ai_analysis_screen(
            main_frame=main_frame,
            build_header=build_header,
            go_back_callback=build_start_screen,
            archive_dir=ARCHIVE_DIR,
            ai_save_dir=AI_REPORTS_DIR
        )
    ).pack(fill="x", pady=6)



    # ======================================================
    # ПРАВЫЙ ИНФО-БЛОК: авто-обновление раз в 10 сек
    # (данные будут подтягиваться из файла tg_status.json через services/tg_counters.py)
    # ======================================================

    # контейнер справа
    info_panel = tk.Frame(main_frame, bg="#f4f6f8")
    info_panel.place(relx=0.72, rely=0.52, anchor="w")

    tk.Label(
        info_panel,
        text="Последние номера проб",
        font=("Segoe UI", 12, "bold"),
        bg="#f4f6f8",
        fg="#111827"
    ).pack(anchor="w", pady=(0, 10))

    # --- строки ---
    swab_row = tk.Frame(info_panel, bg="#f4f6f8")
    swab_row.pack(anchor="w", pady=6)

    air_row = tk.Frame(info_panel, bg="#f4f6f8")
    air_row.pack(anchor="w", pady=6)

    stroke_row = tk.Frame(info_panel, bg="#f4f6f8")  # ✅ МАЗОК
    stroke_row.pack(anchor="w", pady=6)

    # --- переменные текста ---
    swab_left_var = tk.StringVar(value="Номер последнего смыва: —")
    swab_next_var = tk.StringVar(value="—")

    air_left_var = tk.StringVar(value="Номер последней пробы воздуха: —")
    air_next_var = tk.StringVar(value="—")

    stroke_left_var = tk.StringVar(value="Номер последнего мазка: —")  # ✅ МАЗОК
    stroke_next_var = tk.StringVar(value="—")                          # ✅ МАЗОК

    updated_var = tk.StringVar(value="Обновлено: —")

    # --- Смывы: "последний" -> "следующий" ---
    tk.Label(
        swab_row,
        textvariable=swab_left_var,
        font=("Segoe UI", 11),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        swab_row,
        text="  →  ",
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        swab_row,
        textvariable=swab_next_var,
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#16a34a"
    ).pack(side="left")

    # --- Воздух ---
    tk.Label(
        air_row,
        textvariable=air_left_var,
        font=("Segoe UI", 11),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        air_row,
        text="  →  ",
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        air_row,
        textvariable=air_next_var,
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#16a34a"
    ).pack(side="left")

    # --- Мазок ---
    tk.Label(
        stroke_row,
        textvariable=stroke_left_var,
        font=("Segoe UI", 11),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        stroke_row,
        text="  →  ",
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#374151"
    ).pack(side="left")

    tk.Label(
        stroke_row,
        textvariable=stroke_next_var,
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        fg="#16a34a"
    ).pack(side="left")

    tk.Label(
        info_panel,
        textvariable=updated_var,
        font=("Segoe UI", 9),
        bg="#f4f6f8",
        fg="#6b7280"
    ).pack(anchor="w", pady=(10, 0))

    # ✅ хранить after-id, чтобы отменять таймер при уходе со стартового экрана
    refresh_after_id = {"id": None}

    def _apply_numbers(last_swab, last_air, last_stroke):
        if last_swab is None:
            swab_left_var.set("Номер последнего смыва: —")
            swab_next_var.set("—")
        else:
            swab_left_var.set(f"Номер последнего смыва: {last_swab}")
            swab_next_var.set(str(int(last_swab) + 1))

        if last_air is None:
            air_left_var.set("Номер последней пробы воздуха: —")
            air_next_var.set("—")
        else:
            air_left_var.set(f"Номер последней пробы воздуха: {last_air}")
            air_next_var.set(str(int(last_air) + 1))

        if last_stroke is None:
            stroke_left_var.set("Номер последнего мазка: —")
            stroke_next_var.set("—")
        else:
            stroke_left_var.set(f"Номер последнего мазка: {last_stroke}")
            stroke_next_var.set(str(int(last_stroke) + 1))

    def _stop_refresh_timer():
        aid = refresh_after_id.get("id")
        if aid:
            try:
                main_frame.after_cancel(aid)
            except Exception:
                pass
            refresh_after_id["id"] = None

    def _refresh_numbers():
        # если уже ушли с этого экрана — не продолжаем таймер
        if not info_panel.winfo_exists():
            return

        try:
            from services.tg_counters import read_latest_numbers
            last_swab, last_air, last_stroke = read_latest_numbers()
        except Exception:
            last_swab, last_air, last_stroke = None, None, None

        _apply_numbers(last_swab, last_air, last_stroke)

        try:
            import time
            updated_var.set("Обновлено: " + time.strftime("%H:%M:%S"))
        except Exception:
            updated_var.set("Обновлено: —")

        # повтор через 10 секунд (сохраняем id, чтобы отменить)
        refresh_after_id["id"] = main_frame.after(10000, _refresh_numbers)

    # ✅ при уничтожении панели (уход со стартового экрана) — отменяем таймер
    info_panel.bind("<Destroy>", lambda e: _stop_refresh_timer())

    # первый запуск сразу
    _refresh_numbers()

    # ---------- КНОПКА ВЫХОДА (ВНИЗУ СПРАВА) ----------
    bottom_bar = tk.Frame(main_frame, bg="#f4f6f8")
    bottom_bar.pack(side="bottom", fill="x", padx=20, pady=20)

    ttk.Button(
        bottom_bar,
        text="🚪 Выход",
        style="Secondary.TButton",
        command=root.destroy
    ).pack(side="right")






# ======================================================
# СОЗДАНИЕ ОТЧЁТА
# ======================================================

def show_create_report_screen():
    for w in main_frame.winfo_children():
        w.destroy()

    # ---------- ВЕРХНЯЯ ПАНЕЛЬ (НАЗАД) ----------
    top_bar = tk.Frame(main_frame, bg="#f4f6f8")
    top_bar.pack(fill="x", padx=15, pady=(10, 0))

    ttk.Button(
        top_bar,
        text="← Назад",
        style="Secondary.TButton",
        command=build_start_screen
    ).pack(side="left")

    # ---------- КОНТЕНТ ----------
    content = tk.Frame(main_frame, bg="#f4f6f8")
    content.pack(fill="both", expand=True, padx=20, pady=10)

    left = tk.Frame(content, bg="#f4f6f8")
    left.pack(side="left", anchor="n")

    card = tk.Frame(left, bg="white", bd=1, relief="solid")
    card.pack(anchor="nw", padx=(0, 20), pady=(10, 0))

    form = tk.Frame(card, bg="white")
    form.pack(padx=16, pady=16)

    tk.Label(
        form,
        text="Создание отчёта",
        font=("Segoe UI", 16, "bold"),
        bg="white"
    ).pack(anchor="w", pady=(0, 12))

    tk.Label(form, text="Отделение", bg="white").pack(anchor="w")
    dep_var = tk.StringVar()

    from screens.swab_monitoring import get_departments
    dep_cb = ttk.Combobox(
        form,
        textvariable=dep_var,
        values=get_departments(load_config()),
        state="readonly",
        width=35
    )
    dep_cb.pack(fill="x", pady=(4, 14))

    def _refresh_departments_list():
        deps = get_departments(load_config())
        dep_cb.config(values=deps)
        if dep_var.get() and dep_var.get() not in deps:
            dep_var.set("")

    cfg_path = get_config_path()
    try:
        last_mtime = os.path.getmtime(cfg_path)
    except Exception:
        last_mtime = None

    def _poll_config_changes():
        nonlocal last_mtime
        if not dep_cb.winfo_exists():
            return
        try:
            mtime = os.path.getmtime(cfg_path)
        except Exception:
            mtime = last_mtime
        if mtime != last_mtime:
            last_mtime = mtime
            _refresh_departments_list()
        dep_cb.after(2000, _poll_config_changes)

    _poll_config_changes()

    tk.Label(form, text="Период отчёта", bg="white").pack(anchor="w")
    period = tk.Frame(form, bg="white")
    period.pack(anchor="w", pady=(4, 10), fill="x")

    month_var = tk.StringVar()
    year_var = tk.StringVar()
    period_var = tk.StringVar()

    month_names = [
        "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
        "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
    ]

    def set_period_from_date(d):
        month_var.set(month_names[d.month - 1])
        year_var.set(str(d.year))
        period_var.set(f"{month_names[d.month - 1]} {d.year}")

    set_period_from_date(datetime.date.today())

    period_entry = ttk.Entry(
        period,
        textvariable=period_var,
        state="readonly",
        width=22
    )
    period_entry.pack(side="left")

    def open_calendar():
        cal_win = tk.Toplevel(main_frame)
        cal_win.title("Выбор даты")
        cal_win.transient(main_frame)
        cal_win.grab_set()
        cal_win.resizable(False, False)

        body = tk.Frame(cal_win, padx=12, pady=12)
        body.pack(fill="both", expand=True)

        header = tk.Frame(body)
        header.pack(fill="x", pady=(0, 8))

        # текущие значения
        try:
            cur_year = int(year_var.get())
        except Exception:
            cur_year = datetime.date.today().year
        if month_var.get() in month_names:
            cur_month = month_names.index(month_var.get()) + 1
        else:
            cur_month = datetime.date.today().month

        sel_year = {"v": cur_year}
        sel_month = {"v": cur_month}

        def refresh_calendar():
            for w in days_frame.winfo_children():
                w.destroy()

            title_lbl.config(text=f"{month_names[sel_month['v'] - 1]} {sel_year['v']}")

            # дни недели
            hdr = tk.Frame(days_frame)
            hdr.pack(fill="x")
            for d in ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]:
                tk.Label(hdr, text=d, width=3, anchor="center").pack(side="left")

            cal = calendar.Calendar(firstweekday=0)
            for week in cal.monthdayscalendar(sel_year["v"], sel_month["v"]):
                row = tk.Frame(days_frame)
                row.pack(fill="x")
                for day in week:
                    if day == 0:
                        tk.Label(row, text=" ", width=3).pack(side="left")
                    else:
                        btn = tk.Button(
                            row,
                            text=str(day),
                            width=3,
                            relief="flat"
                        )
                        def _pick(d=day):
                            set_period_from_date(datetime.date(sel_year["v"], sel_month["v"], d))
                            cal_win.destroy()
                        btn.config(command=_pick)
                        btn.pack(side="left")

        def prev_month():
            m = sel_month["v"] - 1
            y = sel_year["v"]
            if m < 1:
                m = 12
                y -= 1
            sel_month["v"] = m
            sel_year["v"] = y
            year_spin.delete(0, "end")
            year_spin.insert(0, str(y))
            month_cb.set(month_names[m - 1])
            refresh_calendar()

        def next_month():
            m = sel_month["v"] + 1
            y = sel_year["v"]
            if m > 12:
                m = 1
                y += 1
            sel_month["v"] = m
            sel_year["v"] = y
            year_spin.delete(0, "end")
            year_spin.insert(0, str(y))
            month_cb.set(month_names[m - 1])
            refresh_calendar()

        ttk.Button(header, text="◀", width=3, command=prev_month).pack(side="left")
        title_lbl = tk.Label(header, text="", font=("Segoe UI", 10, "bold"))
        title_lbl.pack(side="left", padx=8)
        ttk.Button(header, text="▶", width=3, command=next_month).pack(side="left")

        controls = tk.Frame(body)
        controls.pack(fill="x", pady=(0, 6))

        month_cb = ttk.Combobox(controls, values=month_names, state="readonly", width=14)
        month_cb.set(month_names[cur_month - 1])
        month_cb.pack(side="left")

        def on_month_change(_=None):
            if month_cb.get() in month_names:
                sel_month["v"] = month_names.index(month_cb.get()) + 1
                refresh_calendar()

        month_cb.bind("<<ComboboxSelected>>", on_month_change)

        year_spin = ttk.Spinbox(controls, from_=2000, to=2100, width=6)
        year_spin.set(str(cur_year))
        year_spin.pack(side="left", padx=(8, 0))

        def on_year_change(_=None):
            try:
                sel_year["v"] = int(year_spin.get())
                refresh_calendar()
            except Exception:
                pass

        year_spin.bind("<FocusOut>", on_year_change)
        year_spin.bind("<Return>", on_year_change)

        days_frame = tk.Frame(body)
        days_frame.pack()

        refresh_calendar()

    ttk.Button(period, text="Выбрать дату", style="Secondary.TButton", command=open_calendar).pack(side="left", padx=(8, 0))

    tk.Label(
        form,
        text="Отчёт формируется за выбранный месяц.",
        bg="white",
        fg="#6b7280"
    ).pack(anchor="w", pady=(0, 12))

    # ---------- ПРОДОЛЖИТЬ ----------
    def proceed():
        global current_department, current_month, current_year

        if not dep_var.get():
            messagebox.showwarning(
                "Не выбрано отделение",
                "Пожалуйста, выберите отделение"
            )
            return

        if not month_var.get() or not year_var.get():
            messagebox.showwarning(
                "Не выбран период",
                "Пожалуйста, выберите месяц и год"
            )
            return

        current_department = dep_var.get()
        current_month = month_var.get()
        current_year = year_var.get()

        # сброс состояния отчёта при новом отчёте
        report_state["microbes"] = None
        report_state["loci"] = None
        report_state["resistance"] = None
        open_microbio_files_wizard(main_frame, on_done=show_report_workspace)

    ttk.Button(
        form,
        text="Продолжить",
        style="Main.TButton",
        command=proceed
    ).pack(anchor="w", pady=(6, 0))


def open_microbio_files_wizard(parent, on_done):
    from analysis.microbes import analyze_microbes
    from analysis.loci import analyze_loci
    from analysis.resistance import analyze_resistance

    win = tk.Toplevel(parent)
    win.title("Файлы для отчёта")
    win.transient(parent)
    win.grab_set()
    win.resizable(False, False)

    body = tk.Frame(win, padx=18, pady=16, bg="white")
    body.pack(fill="both", expand=True)

    tk.Label(
        body,
        text="Подготовка отчёта по микробиологическому мониторингу",
        font=("Segoe UI", 13, "bold"),
        bg="white",
        fg="#111827",
    ).pack(anchor="w")

    tk.Label(
        body,
        text=(
            "Для построения отчёта нужно добавить 3 Excel‑файла:\n"
            "• Микроорганизмы\n"
            "• Локусы\n"
            "• Резистентность\n\n"
            "Файлы можно прикреплять по одному — система проверит формат\n"
            "и автоматически определит тип каждого файла."
        ),
        font=("Segoe UI", 10),
        bg="white",
        fg="#6b7280",
        justify="left",
    ).pack(anchor="w", pady=(6, 12))

    kinds = [
        ("microbes", "Микроорганизмы"),
        ("loci", "Локусы"),
        ("resistance", "Резистентность"),
    ]
    files = {}
    status_vars = {}
    rows = {}

    create_btn_visible = {"v": False}

    def _update_ready():
        ready = all(report_state.get(k) for k, _ in kinds)
        create_btn.config(state="normal" if ready else "disabled")
        if ready and not create_btn_visible["v"]:
            create_btn.pack(anchor="e")
            create_btn_visible["v"] = True
        elif not ready and create_btn_visible["v"]:
            create_btn.pack_forget()
            create_btn_visible["v"] = False

    def _set_status(kind: str, text: str, ok: bool = False):
        v = status_vars[kind]
        v.set(text)
        rows[kind]["status"].config(fg="#16a34a" if ok else "#6b7280")

    def _apply_microbes_result(result):
        from utils.gram import classify_gram
        for m in result.get("microbes", []):
            if m.get("gram_source") == "manual":
                continue
            gram_raw = (m.get("gram") or "").strip()
            if gram_raw in ("Gram+", "Gram-", "Грибы", "Не классифицировано"):
                m["gram"] = gram_raw
            else:
                m["gram"] = normalize_gram(classify_gram(m.get("microbe", "")))
            m["gram_source"] = m.get("gram_source") or "auto"
        report_state["microbes"] = result

    def _assign(kind: str, file_path: str, result):
        if kind == "microbes":
            _apply_microbes_result(result)
        else:
            report_state[kind] = result
        files[kind] = file_path
        _set_status(kind, f"✓ {os.path.basename(file_path)}", ok=True)
        _update_ready()

    def _try_parse_kind(kind: str, file_path: str):
        try:
            if kind == "microbes":
                return analyze_microbes(file_path), None
            if kind == "loci":
                return analyze_loci(file_path), None
            return analyze_resistance(file_path), None
        except Exception as e:
            return None, str(e)

    def _detect_and_parse(file_path: str):
        attempts = {}
        errors = {}

        for kind, _ in kinds:
            result, err = _try_parse_kind(kind, file_path)
            if result is not None:
                attempts[kind] = result
            else:
                errors[kind] = err or "Не удалось прочитать файл"

        if len(attempts) == 1:
            kind = list(attempts.keys())[0]
            if report_state.get(kind):
                if not messagebox.askyesno("Заменить файл?", f"Файл для «{dict(kinds)[kind]}» уже выбран. Заменить?"):
                    return
            _assign(kind, file_path, attempts[kind])
            return

        if len(attempts) == 0:
            msg = "Не удалось определить тип файла.\n\n"
            msg += "Проверьте, что это правильные выгрузки:\n"
            msg += "• Микроорганизмы: столбцы «Обнар. микроорг.» и «COUNT(*)»\n"
            msg += "• Локусы: столбцы «Локус» и «COUNT(*)»\n"
            msg += "• Резистентность: столбцы с R/S/I и микроорганизмами\n\n"
            msg += "Детали:\n"
            for k, err in errors.items():
                msg += f"- {dict(kinds)[k]}: {err}\n"
            messagebox.showerror("Файл не распознан", msg)
            return

        # несколько вариантов — спросить пользователя
        kind = _ask_kind(win, [k for k in attempts.keys()])
        if not kind:
            return
        if report_state.get(kind):
            if not messagebox.askyesno("Заменить файл?", f"Файл для «{dict(kinds)[kind]}» уже выбран. Заменить?"):
                return
        _assign(kind, file_path, attempts[kind])

    def _ask_kind(parent_win, options):
        top = tk.Toplevel(parent_win)
        top.title("Тип файла")
        top.transient(parent_win)
        top.grab_set()
        top.resizable(False, False)

        frm = tk.Frame(top, padx=14, pady=12)
        frm.pack(fill="both", expand=True)

        tk.Label(frm, text="Уточните тип файла:", font=("Segoe UI", 11, "bold")).pack(anchor="w", pady=(0, 8))
        var = tk.StringVar(value=options[0])
        for k in options:
            ttk.Radiobutton(frm, text=dict(kinds)[k], value=k, variable=var).pack(anchor="w")

        res = {"val": None}

        def _ok():
            res["val"] = var.get()
            top.destroy()

        def _cancel():
            top.destroy()

        btns = tk.Frame(frm)
        btns.pack(fill="x", pady=(8, 0))
        ttk.Button(btns, text="Отмена", style="Secondary.TButton", command=_cancel).pack(side="right")
        ttk.Button(btns, text="Выбрать", style="Main.TButton", command=_ok).pack(side="right", padx=(0, 6))

        top.wait_window()
        return res["val"]

    def add_files():
        paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx *.xls")])
        if not paths:
            return
        for p in paths:
            _detect_and_parse(p)

    list_box = tk.Frame(body, bg="white")
    list_box.pack(fill="x", pady=(0, 12))

    for kind, title in kinds:
        row = tk.Frame(list_box, bg="white")
        row.pack(fill="x", pady=4)

        tk.Label(row, text=title, width=18, anchor="w", bg="white").pack(side="left")
        status_var = tk.StringVar(value="Файл не выбран")
        status_lbl = tk.Label(row, textvariable=status_var, bg="white", fg="#6b7280", anchor="w")
        status_lbl.pack(side="left", fill="x", expand=True, padx=(6, 0))

        def _choose_for_kind(k=kind):
            path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
            if not path:
                return
            result, err = _try_parse_kind(k, path)
            if result is None:
                messagebox.showerror(
                    "Файл не подходит",
                    f"Файл не подходит для «{dict(kinds)[k]}».\n\n{err}"
                )
                return
            if report_state.get(k):
                if not messagebox.askyesno("Заменить файл?", f"Файл для «{dict(kinds)[k]}» уже выбран. Заменить?"):
                    return
            _assign(k, path, result)

        ttk.Button(row, text="Выбрать", style="Secondary.TButton", command=_choose_for_kind).pack(side="right")

        status_vars[kind] = status_var
        rows[kind] = {"status": status_lbl}

    add_row = tk.Frame(body, bg="white")
    add_row.pack(fill="x", pady=(0, 10))
    ttk.Button(add_row, text="Добавить файлы", style="Main.TButton", command=add_files).pack(side="left")

    create_btn = ttk.Button(body, text="Создать отчёт", style="Main.TButton", state="disabled")

    def _finish():
        win.destroy()
        on_done()

    create_btn.config(command=_finish)

    _update_ready()



# ======================================================
# HEADER
# ======================================================

def open_settings_dialog(parent):
    cfg = load_config()
    webdav_url = (cfg.get("webdav_url") or "").strip()

    win = tk.Toplevel(parent)
    win.title("Настройки")
    win.transient(parent)
    win.grab_set()
    win.resizable(False, False)

    frame = tk.Frame(win, padx=16, pady=14)
    frame.pack(fill="both", expand=True)

    tk.Label(
        frame,
        text="Настройки",
        font=("Segoe UI", 12, "bold")
    ).pack(anchor="w", pady=(0, 10))

    # --- Журнал ---
    tk.Label(frame, text="Excel-журнал смывов:", font=("Segoe UI", 9, "bold")).pack(anchor="w")
    journal_lbl = tk.Label(
        frame,
        text=cfg.get("swabs_journal_xlsx", ""),
        wraplength=520,
        justify="left"
    )
    journal_lbl.pack(anchor="w", pady=(0, 8))

    btns_row = tk.Frame(frame)
    btns_row.pack(anchor="w", pady=(0, 14))

    def set_journal_path():
        new_file = filedialog.askopenfilename(
            title="Выберите Excel-журнал",
            filetypes=[("Excel", "*.xlsx *.xls")]
        )
        if not new_file:
            return
        cfg["swabs_journal_xlsx"] = new_file
        save_config(cfg)
        journal_lbl.config(text=new_file)

    def show_reload_hint():
        messagebox.showinfo(
            "Обновить",
            "Откройте «Мониторинг смывов» и нажмите «Обновить», чтобы перечитать журнал."
        )

    ttk.Button(btns_row, text="Путь файла", style="Secondary.TButton", command=set_journal_path).pack(side="left", padx=(0, 6))
    ttk.Button(btns_row, text="Обновить", style="Secondary.TButton", command=show_reload_hint).pack(side="left")

    # --- Шаблон отчёта ---
    tk.Label(frame, text="Шаблон отчёта (DOCX):", font=("Segoe UI", 9, "bold")).pack(anchor="w")
    tpl_lbl = tk.Label(
        frame,
        text=cfg.get("swabs_template_path", ""),
        wraplength=520,
        justify="left"
    )
    tpl_lbl.pack(anchor="w", pady=(0, 8))

    def set_template_path():
        new_file = filedialog.askopenfilename(
            title="Выберите шаблон отчёта",
            filetypes=[("Word", "*.docx")]
        )
        if not new_file:
            return
        if (cfg.get("webdav_url") or "").strip():
            try:
                target = os.path.join(DATA_ROOT, "documents", "шаблон.docx")
                os.makedirs(os.path.dirname(target), exist_ok=True)
                shutil.copy2(new_file, target)
                webdav_sync.upload_file(target, DATA_ROOT)
                new_file = target
            except Exception:
                pass
        cfg["swabs_template_path"] = new_file
        save_config(cfg)
        tpl_lbl.config(text=new_file)

    ttk.Button(frame, text="Изменить шаблон", style="Secondary.TButton", command=set_template_path).pack(anchor="w", pady=(0, 14))

    # --- Папка сохранения отчётов ---
    tk.Label(frame, text="Папка сохранения отчётов:", font=("Segoe UI", 9, "bold")).pack(anchor="w")
    out_lbl = tk.Label(
        frame,
        text=cfg.get("swabs_report_dir", ""),
        wraplength=520,
        justify="left"
    )
    out_lbl.pack(anchor="w", pady=(0, 8))

    def set_report_dir():
        new_dir = filedialog.askdirectory(title="Выберите папку для сохранения отчётов")
        if not new_dir:
            return
        cfg["swabs_report_dir"] = new_dir
        save_config(cfg)
        out_lbl.config(text=new_dir)

    ttk.Button(frame, text="Изменить папку", style="Secondary.TButton", command=set_report_dir).pack(anchor="w")

    # --- Привязки отделений (aliases) ---
    def _aliases_path_for_settings() -> str:
        p = (cfg.get("swabs_aliases_json") or "").strip()
        if webdav_url:
            return os.path.join(DATA_ROOT, "config", "swabs_dep_aliases.json")
        if not p:
            p = os.path.join(DATA_ROOT, "config", "swabs_dep_aliases.json")
            cfg["swabs_aliases_json"] = p
            save_config(cfg)
        return p

    aliases_path = _aliases_path_for_settings()
    try:
        from analysis.dep_mapper import save_aliases
        os.makedirs(os.path.dirname(aliases_path), exist_ok=True)
        if not os.path.exists(aliases_path):
            save_aliases(aliases_path, {})
            if webdav_url:
                webdav_sync.upload_file(aliases_path, DATA_ROOT)
    except Exception:
        pass

    tk.Label(frame, text="Привязки отделений (aliases):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(14, 0))
    alias_lbl = tk.Label(
        frame,
        text=aliases_path,
        wraplength=520,
        justify="left"
    )
    alias_lbl.pack(anchor="w", pady=(0, 8))

    alias_btns = tk.Frame(frame)
    alias_btns.pack(anchor="w")

    def edit_aliases_dialog():
        from analysis.dep_mapper import load_aliases, save_aliases
        from screens.dep_map_dialog import ask_user_map_unknowns
        from screens.swab_monitoring import get_departments

        if webdav_url:
            try:
                webdav_sync.sync_down(DATA_ROOT)
            except Exception:
                pass

        aliases = load_aliases(aliases_path)
        unknown = sorted(aliases.keys())
        if not unknown:
            messagebox.showinfo("Привязки", "Список привязок пуст.")
            return

        departments = get_departments(cfg)
        updated = ask_user_map_unknowns(win, unknown, departments, preset=aliases)
        if updated is None:
            return

        merged = dict(aliases)
        merged.update(updated)
        try:
            save_aliases(aliases_path, merged)
            if webdav_url:
                webdav_sync.upload_file(aliases_path, DATA_ROOT)
            messagebox.showinfo("Привязки", "Сохранено.")
        except Exception as e:
            messagebox.showerror("Привязки", f"Не удалось сохранить:\n{e}")

    ttk.Button(alias_btns, text="Редактировать", style="Secondary.TButton", command=edit_aliases_dialog).pack(side="left")

    # --- База отделений ---
    tk.Label(frame, text="База отделений:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(14, 0))
    list_frame = tk.Frame(frame)
    list_frame.pack(fill="x", pady=(4, 8))

    deps_list = tk.Listbox(list_frame, width=60, height=8)
    deps_scroll = ttk.Scrollbar(list_frame, orient="vertical", command=deps_list.yview)
    deps_list.configure(yscrollcommand=deps_scroll.set)
    deps_list.pack(side="left", fill="both", expand=True)
    deps_scroll.pack(side="right", fill="y")

    from screens.swab_monitoring import get_departments
    for d in get_departments(cfg):
        deps_list.insert(tk.END, d)

    entry_row = tk.Frame(frame)
    entry_row.pack(fill="x", pady=(0, 6))

    dep_var = tk.StringVar()
    dep_entry = ttk.Entry(entry_row, textvariable=dep_var)
    dep_entry.pack(side="left", fill="x", expand=True)

    def _save_departments():
        deps = [deps_list.get(i) for i in range(deps_list.size())]
        deps = [d.strip() for d in deps if d and str(d).strip()]
        cfg["departments"] = deps
        save_config(cfg)

    def add_department():
        val = (dep_var.get() or "").strip()
        if not val:
            return
        existing = [deps_list.get(i) for i in range(deps_list.size())]
        if val in existing:
            return
        deps_list.insert(tk.END, val)
        dep_var.set("")
        _save_departments()

    def remove_department():
        sel = deps_list.curselection()
        if not sel:
            return
        for idx in reversed(sel):
            deps_list.delete(idx)
        _save_departments()

    ttk.Button(entry_row, text="Добавить", style="Secondary.TButton", command=add_department).pack(side="left", padx=(6, 0))
    ttk.Button(entry_row, text="Удалить выбранные", style="Secondary.TButton", command=remove_department).pack(side="left", padx=(6, 0))

    tk.Label(
        frame,
        text="Изменения применятся при следующем открытии «Мониторинг смывов».",
        fg="#6b7280"
    ).pack(anchor="w", pady=(0, 4))

def build_header(parent, back_callback=None):
    header = tk.Frame(parent, bg="#e5e7eb")
    header.pack(fill="x")

    ttk.Button(
        header,
        text="← Назад" if back_callback else "← На главную",
        style="Secondary.TButton",
        command=back_callback or build_start_screen
    ).pack(side="left", padx=15, pady=8)

    right = tk.Frame(header, bg="#e5e7eb")
    right.pack(side="right", padx=15)

    ttk.Button(
        right,
        text="⚙ Настройки",
        style="Secondary.TButton",
        command=lambda: open_settings_dialog(parent)
    ).pack(side="right")

    tk.Label(
        right,
        text="ЭпидМонитор",
        font=("Segoe UI", 11, "bold"),
        bg="#e5e7eb"
    ).pack(side="right", padx=(0, 10))


# ======================================================
# РАБОЧЕЕ ПРОСТРАНСТВО
# ======================================================

def show_report_workspace():
    for w in main_frame.winfo_children():
        w.destroy()

    build_header(main_frame)

    # ---------- NOTEBOOK ----------
    notebook = ttk.Notebook(main_frame)
    notebook.pack(expand=True, fill="both", padx=10, pady=(10, 0))

    tab_microbes = tk.Frame(notebook, bg="white")
    notebook.add(tab_microbes, text="Микроорганизмы")
    build_microbes_tab(tab_microbes, report_state, normalize_gram)

    tab_loci = tk.Frame(notebook, bg="white")
    notebook.add(tab_loci, text="Локусы")
    build_loci_tab(tab_loci, report_state)

    tab_res = tk.Frame(notebook, bg="white")
    notebook.add(tab_res, text="Резистентность")
    build_resistance_tab(tab_res, report_state)

    # ---------- НИЖНЯЯ ПАНЕЛЬ (КНОПКА) ----------
    bottom_bar = tk.Frame(main_frame, bg="#f4f6f8")
    bottom_bar.pack(fill="x", pady=10)

    ttk.Button(
        bottom_bar,
        text="💾 Сохранить отчёт",
        style="Main.TButton",
        command=save_report
    ).pack(pady=5)



# ======================================================
# МИКРООРГАНИЗМЫ + GRAM
# ======================================================

# ======================================================
# ARCHIVE + OUTLOOK
# ======================================================

import time
import win32com.client
import win32gui
import win32con


# ------------------------------------------------------
# ОТПРАВКА В OUTLOOK
# ------------------------------------------------------

def send_to_outlook(filepath):
    if not os.path.exists(filepath):
        messagebox.showerror("Ошибка", "Файл не найден")
        return

    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)

        mail.Subject = "Эпидемиологический отчёт"
        mail.Body = (
            "Здравствуйте!\n\n"
            "Направляю эпидемиологический отчёт.\n\n"
            "С уважением."
        )

        mail.Attachments.Add(os.path.abspath(filepath))
        mail.Display()

        # даём Outlook создать окно
        time.sleep(0.4)

        inspector = mail.GetInspector
        inspector.Activate()

        time.sleep(0.2)

        # поднимаем окно Outlook поверх
        def enum_windows(hwnd, result):
            if win32gui.IsWindowVisible(hwnd):
                title = win32gui.GetWindowText(hwnd)
                if "Outlook" in title:
                    result.append(hwnd)

        windows = []
        win32gui.EnumWindows(enum_windows, windows)

        if windows:
            hwnd = windows[0]
            win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
            win32gui.SetForegroundWindow(hwnd)

    except Exception as e:
        messagebox.showerror("Ошибка Outlook", str(e))

# ------------------------------------------------------
# ARCHIVE HELPERS
# ------------------------------------------------------

def _build_scroll_area(container, bg="#f4f6f8"):
    canvas = tk.Canvas(container, bg=bg, highlightthickness=0)
    scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
    scroll_frame = tk.Frame(canvas, bg=bg)

    scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    return scroll_frame


def _bind_row_hover(row, bg_normal="#f4f6f8", bg_hover="#e5e7eb"):
    def on_enter(event):
        row.configure(bg=bg_hover)
        for c in row.winfo_children():
            c.configure(bg=bg_hover)

    def on_leave(event):
        row.configure(bg=bg_normal)
        for c in row.winfo_children():
            c.configure(bg=bg_normal)

    row.bind("<Enter>", on_enter)
    row.bind("<Leave>", on_leave)
    for c in row.winfo_children():
        c.bind("<Enter>", on_enter)
        c.bind("<Leave>", on_leave)


def _add_archive_file_row(parent, display_name, full_path, on_delete, *, prefix="", cursor=None):
    row = tk.Frame(parent, bg="#f4f6f8")
    row.pack(fill="x", padx=12, pady=3)

    label = tk.Label(
        row,
        text=f"{prefix}{display_name}" if prefix else display_name,
        bg="#f4f6f8",
        anchor="w",
        font=("Segoe UI", 10),
        cursor=cursor
    )
    label.pack(fill="x", padx=4, pady=2)

    _bind_row_hover(row)

    row.bind("<Double-Button-1>", lambda e, p=full_path: os.startfile(p))
    label.bind("<Double-Button-1>", lambda e, p=full_path: os.startfile(p))

    menu = tk.Menu(row, tearoff=0)
    menu.add_command(label="✉ Отправить", command=lambda p=full_path: send_to_outlook(p))
    menu.add_separator()
    def _delete_and_sync(p=full_path):
        os.remove(p)
        webdav_sync.delete_path(p, DATA_ROOT)
        on_delete()

    menu.add_command(label="🗑 Удалить", command=_delete_and_sync)

    def show_menu(event, m=menu):
        m.tk_popup(event.x_root, event.y_root)

    row.bind("<Button-3>", show_menu)
    label.bind("<Button-3>", show_menu)

# ------------------------------------------------------
# ВЫБОР АРХИВА
# ------------------------------------------------------

def build_archive_choice_screen():
    for w in main_frame.winfo_children():
        w.destroy()

    build_header(main_frame, back_callback=build_start_screen)


    tk.Label(
        main_frame,
        text="Архив",
        font=("Segoe UI", 18, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(40, 30))

    cards = tk.Frame(main_frame, bg="#f4f6f8")
    cards.pack()

    # ---------- МИКРОБИОЛОГИЯ ----------
    ttk.Button(
        cards,
        text="🧫 Архив микробиологического мониторинга",
        style="Main.TButton",
        command=build_microbio_archive_screen
    ).pack(fill="x", ipadx=20, ipady=10, pady=10)

    # ---------- СМЫВЫ ----------
    ttk.Button(
        cards,
        text="🧼 Архив смывов",
        style="Main.TButton",
        command=build_swab_archive_screen
    ).pack(fill="x", ipadx=20, ipady=10)


# ------------------------------------------------------
# АРХИВ СМЫВОВ (ТОЛЬКО ПО ОТДЕЛЕНИЯМ)
# ------------------------------------------------------

def build_swab_archive_screen():
    import datetime as dt
    import tkinter as tk
    from tkinter import ttk

    for w in main_frame.winfo_children():
        w.destroy()

    build_header(main_frame, back_callback=build_archive_choice_screen)

    tk.Label(
        main_frame,
        text="Архив смывов",
        font=("Segoe UI", 18, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(25, 10))

    container = tk.Frame(main_frame, bg="#f4f6f8")
    container.pack(expand=True, fill="both", padx=20, pady=10)

    # --- если папка недоступна ---
    if not os.path.exists(SWABS_ARCHIVE_DIR):
        tk.Label(
            container,
            text=f"Папка архива недоступна:\n{SWABS_ARCHIVE_DIR}",
            font=("Segoe UI", 12),
            bg="#f4f6f8"
        ).pack(pady=40)
        return

    # --- скролл ---
    scroll_frame = _build_scroll_area(container)

    # --- собираем docx только из одной папки (без подпапок) ---
    all_files = [
        f for f in os.listdir(SWABS_ARCHIVE_DIR)
        if os.path.isfile(os.path.join(SWABS_ARCHIVE_DIR, f)) and f.lower().endswith(".docx")
    ]

    if not all_files:
        tk.Label(
            scroll_frame,
            text="Нет отчётов (.docx) в папке архива",
            font=("Segoe UI", 12),
            bg="#f4f6f8"
        ).pack(pady=40)
        return

    # --- сортировка: сначала пробуем дату dd.mm.yyyy из имени, иначе по mtime ---
    date_re = re.compile(r"(\d{2}\.\d{2}\.\d{4})")

    def sort_key(fname: str):
        m = date_re.search(fname)
        if m:
            s = m.group(1)
            try:
                d = dt.datetime.strptime(s, "%d.%m.%Y")
                return (0, d, fname.lower())
            except Exception:
                pass

        full = os.path.join(SWABS_ARCHIVE_DIR, fname)
        try:
            ts = os.path.getmtime(full)
        except Exception:
            ts = 0
        return (1, dt.datetime.fromtimestamp(ts), fname.lower())

    all_files.sort(key=sort_key, reverse=True)

    # --- список файлов ---
    box = tk.LabelFrame(
        scroll_frame,
        text="Отчёты (по дате)",
        font=("Segoe UI", 11, "bold"),
        bg="#f4f6f8",
        padx=10,
        pady=8
    )
    box.pack(fill="x", padx=10, pady=10)

    for file in all_files:
        full_path = os.path.join(SWABS_ARCHIVE_DIR, file)
        _add_archive_file_row(
            box,
            file,
            full_path,
            build_swab_archive_screen,
            prefix="📄 ",
            cursor="hand2"
        )




# ------------------------------------------------------
# АРХИВ МИКРОБИОЛОГИЧЕСКОГО МОНИТОРИНГА
# ------------------------------------------------------

def build_microbio_archive_screen():
    for w in main_frame.winfo_children():
        w.destroy()

    build_header(main_frame, back_callback=build_archive_choice_screen)
    webdav_sync.sync_down(DATA_ROOT)


    tk.Label(
        main_frame,
        text="Архив микробиологического мониторинга",
        font=("Segoe UI", 18, "bold"),
        bg="#f4f6f8"
    ).pack(pady=(25, 10))

    container = tk.Frame(main_frame, bg="#f4f6f8")
    container.pack(expand=True, fill="both", padx=20, pady=10)

    departments = sorted(
        d for d in os.listdir(ARCHIVE_DIR)
        if os.path.isdir(os.path.join(ARCHIVE_DIR, d))
    )

    if not departments:
        tk.Label(
            container,
            text="Архив пуст",
            font=("Segoe UI", 12),
            bg="#f4f6f8"
        ).pack(pady=40)
        return

    scroll_frame = _build_scroll_area(container)

    # ---------- ОТДЕЛЕНИЯ ----------
    for dep in departments:
        dep_path = os.path.join(ARCHIVE_DIR, dep)
        files = sorted(
            [f for f in os.listdir(dep_path) if f.endswith(".docx")],
            reverse=True
        )

        box = tk.LabelFrame(
            scroll_frame,
            text=dep,
            font=("Segoe UI", 11, "bold"),
            bg="#f4f6f8",
            padx=10,
            pady=8
        )
        box.pack(fill="x", padx=10, pady=10)

        if not files:
            tk.Label(
                box,
                text="Нет отчётов",
                bg="#f4f6f8",
                fg="#6b7280"
            ).pack(anchor="w")
            continue

        # ---------- ФАЙЛЫ ----------
        for filename in files:
            full_path = os.path.join(dep_path, filename)
            _add_archive_file_row(
                box,
                filename,
                full_path,
                build_microbio_archive_screen
            )



# ======================================================
# СТАРТ ПРИЛОЖЕНИЯ (ОДИН РАЗ)
# ======================================================

def start_app():
    """Запуск настоящего окна после заставки."""
    global splash_window
    try:
        if splash_window is not None:
            splash_window.destroy()
    except Exception:
        pass

    root.deiconify()       # показать главное окно
    root.state("zoomed")   # развернуть на весь экран
    build_login_screen()   # переход к логину

if __name__ == "__main__":
    show_splash_screen()

    def delayed_start():
        root.after(4000, start_app)

    root.after_idle(delayed_start)
    root.mainloop()





