from __future__ import annotations

import os
import re
import time
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Optional, Tuple

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# =========================
# АВТО-СОХРАНЕНИЕ (СЕТЕВЫЕ ПАПКИ)
# =========================
AUTO_WORD_DIR = r"\\w00164\e$\ДОКУМЕНТЫ ОТДЕЛА\2026\ИСМП 2026\Протоколы исследований\Автоматические отчеты\Word"
AUTO_PDF_DIR = r"\\w00164\e$\ДОКУМЕНТЫ ОТДЕЛА\2026\ИСМП 2026\Протоколы исследований\Автоматические отчеты\PDF"


# -------------------------
# helpers
# -------------------------

def _fmt_date_ddmmyyyy(day: pd.Timestamp) -> str:
    return pd.Timestamp(day).strftime("%d.%m.%Y")


def _norm_text(x: object) -> str:
    if x is None:
        return ""
    s = str(x).strip()
    if s in ("", "—", "–"):
        return ""
    return s


def _norm_culture(x: object) -> str:
    s = _norm_text(x)
    if s in ("", "-"):
        return "-"
    return s


def _find_col_like(df: pd.DataFrame, needle: str) -> Optional[str]:
    needle = (needle or "").strip().lower()
    for c in df.columns:
        if needle in str(c).strip().lower():
            return c
    return None


def _set_run_font(run, *, name: str = "Times New Roman", size_pt: int = 12, bold: Optional[bool] = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold

    # фикс кириллицы/подмены шрифта Word'ом
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _set_cell_text(cell, text: str, *, bold: bool = False,
                   font_name: str = "Times New Roman", font_size: int = 12) -> None:
    cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(str(text))
    _set_run_font(run, name=font_name, size_pt=font_size, bold=bold)


def _replace_in_paragraph(paragraph, pattern: str, repl: str) -> bool:
    full = "".join(r.text for r in paragraph.runs)
    new = re.sub(pattern, repl, full)
    if new == full:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new)
    return True


def _set_paragraph_text(paragraph, text: str, *, bold: Optional[bool] = None,
                        font_name: str = "Times New Roman", font_size: int = 12) -> None:
    for r in paragraph.runs:
        r.text = ""
    run = paragraph.add_run(text)
    _set_run_font(run, name=font_name, size_pt=font_size, bold=bold)


def _enforce_landscape(doc: Document) -> None:
    """
    Всегда альбомная ориентация (горизонтальная) во всех секциях.
    """
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def _set_table_borders(table) -> None:
    """
    Явные границы таблицы (окантовка + внутренние линии).
    """
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.append(tblPr)

    tblBorders = OxmlElement("w:tblBorders")

    def _border(tag: str):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        return el

    tblBorders.append(_border("w:top"))
    tblBorders.append(_border("w:left"))
    tblBorders.append(_border("w:bottom"))
    tblBorders.append(_border("w:right"))
    tblBorders.append(_border("w:insideH"))
    tblBorders.append(_border("w:insideV"))

    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _clear_table_data_rows(table) -> None:
    """
    Удаляем все строки кроме заголовка (первая строка).
    """
    while len(table.rows) > 1:
        tr = table.rows[1]._tr
        tr.getparent().remove(tr)


# -------------------------
# Word -> PDF
# -------------------------

def save_pdf_via_word(docx_path: str, pdf_path: str) -> None:
    """
    Конвертирует DOCX в PDF через установленный Microsoft Word (Windows only).
    """
    import win32com.client  # type: ignore

    docx_path = os.path.normpath(os.path.abspath(docx_path))
    pdf_path = os.path.normpath(os.path.abspath(pdf_path))

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX не найден: {docx_path}")

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    doc = None
    try:
        time.sleep(0.1)
        doc = word.Documents.Open(docx_path, ReadOnly=1)
        # 17 = wdFormatPDF
        doc.SaveAs(pdf_path, FileFormat=17)
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        word.Quit()


# -------------------------
# header fill
# -------------------------

def _set_sampling_date_and_dep(doc: Document, dep: str, date_str: str) -> None:
    """
    Для КАЖДОГО блока:
    1) заменяем дату в строках "Дата отбора проб: ..."
    2) следующий непустой абзац после этой строки считаем отделением:
       жирный, Times New Roman 14
    """
    paras = doc.paragraphs

    for i, p in enumerate(paras):
        replaced = _replace_in_paragraph(
            p,
            r"(Дата\s+отбора\s+проб:\s*)(.+)$",
            r"\g<1>" + date_str
        )
        if not replaced:
            continue

        j = i + 1
        while j < len(paras) and not (paras[j].text or "").strip():
            j += 1
        if j < len(paras):
            _set_paragraph_text(
                paras[j],
                dep,
                bold=True,
                font_name="Times New Roman",
                font_size=14
            )


# -------------------------
# tables fill
# -------------------------

def _fill_swabs_table(table, df_swabs: pd.DataFrame) -> None:
    """
    Таблица смывов:
    № | Наименование помещения | Место отбора проб | Выделенная культура (эпид.значима)
    """
    room_col = _find_col_like(df_swabs, "Наименование помещения")
    place_col = _find_col_like(df_swabs, "Место отбора проб")
    cult_col = _find_col_like(df_swabs, "Выделенная культура")

    if not room_col or not place_col or not cult_col:
        return

    _set_table_borders(table)
    _clear_table_data_rows(table)

    n = 1
    for _, row in df_swabs.iterrows():
        room = _norm_text(row.get(room_col, ""))
        place = _norm_text(row.get(place_col, ""))
        cult = _norm_culture(row.get(cult_col, ""))

        if not room and not place and cult == "-":
            continue

        cells = table.add_row().cells
        _set_cell_text(cells[0], str(n), font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[1], room, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[2], place, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[3], cult, font_name="Times New Roman", font_size=12)
        n += 1


def _fill_air_table(table, df_air: pd.DataFrame) -> None:
    """
    Таблица воздуха:
    № | Наименование помещения | Условия отбора (...) | ОМЧ | Выделенная культура (эпид.значима)
    """
    room_col = _find_col_like(df_air, "Наименование помещения")
    cond_col = _find_col_like(df_air, "Условия отбора")
    omch_col = _find_col_like(df_air, "ОМЧ")
    cult_col = _find_col_like(df_air, "Выделенная культура")

    if not room_col or not cond_col or not omch_col or not cult_col:
        return

    _set_table_borders(table)
    _clear_table_data_rows(table)

    n = 1
    for _, row in df_air.iterrows():
        room = _norm_text(row.get(room_col, ""))
        cond = _norm_text(row.get(cond_col, ""))
        omch = _norm_text(row.get(omch_col, ""))
        cult = _norm_culture(row.get(cult_col, ""))

        if not room and not cond and not omch and cult == "-":
            continue

        cells = table.add_row().cells
        _set_cell_text(cells[0], str(n), font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[1], room, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[2], cond, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[3], omch, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[4], cult, font_name="Times New Roman", font_size=12)
        n += 1


def _fill_smears_table(table, df_smears: pd.DataFrame) -> None:
    """
    Таблица мазков (персонал):
    № | ФИО сотрудника | Место отбора проб | Выделенная культура (эпид.значима)
    """
    fio_col = _find_col_like(df_smears, "ФИО")
    place_col = _find_col_like(df_smears, "Место отбора проб")
    cult_col = _find_col_like(df_smears, "Выделенная культура")

    if not fio_col or not place_col or not cult_col:
        return

    _set_table_borders(table)
    _clear_table_data_rows(table)

    n = 1
    for _, row in df_smears.iterrows():
        fio = _norm_text(row.get(fio_col, ""))
        place = _norm_text(row.get(place_col, ""))
        cult = _norm_culture(row.get(cult_col, ""))

        if not fio and not place and cult == "-":
            continue

        cells = table.add_row().cells
        _set_cell_text(cells[0], str(n), font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[1], fio, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[2], place, font_name="Times New Roman", font_size=12)
        _set_cell_text(cells[3], cult, font_name="Times New Roman", font_size=12)
        n += 1


# -------------------------
# section removal + page breaks (clean)
# -------------------------

def _insert_page_break_at_body_index(doc: Document, idx: int) -> None:
    """
    Вставляет разрыв страницы в body перед элементом с индексом idx.
    На уровне XML, чтобы разрыв гарантированно встал до заголовков акта.
    """
    body = doc._element.body  # pylint: disable=protected-access
    children = list(body.iterchildren())
    if idx <= 0 or idx > len(children):
        return

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)

    body.insert(idx, p)


def _is_empty_paragraph_el(p_el) -> bool:
    """
    True если это реально пустой абзац, который можно удалить:
    - нет текста
    - нет картинок/рисунков
    - это не pagebreak-параграф
    """
    if p_el.tag != qn("w:p"):
        return False

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # не трогаем абзацы с разрывом страницы
    br = p_el.find(".//w:br", namespaces=ns)
    if br is not None and br.get(qn("w:type")) == "page":
        return False

    # не трогаем абзацы с рисунками
    if p_el.find(".//w:drawing", namespaces=ns) is not None:
        return False

    # текст
    texts = []
    for t in p_el.findall(".//w:t", namespaces=ns):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip() == ""


def _start_each_act_on_new_page(doc: Document) -> None:
    """
    Каждый акт с новой страницы, БЕЗ огромных отступов.
    Удаляем пустые абзацы в начале каждого акта и ставим break перед первым нормальным абзацем.
    """
    body = doc._element.body  # pylint: disable=protected-access

    children = list(body.iterchildren())
    tbl_els = [t._element for t in doc.tables]

    tbl_pos: list[int] = []
    for tbl_el in tbl_els:
        for i, el in enumerate(children):
            if el is tbl_el:
                tbl_pos.append(i)
                break

    if len(tbl_pos) <= 1:
        return

    # идём с конца: безопасно для индексов
    for k in range(len(tbl_pos) - 1, 0, -1):
        start_idx = tbl_pos[k - 1] + 1

        children = list(body.iterchildren())

        # вычищаем пустые абзацы, которые дают "дыру" сверху
        while start_idx < len(children) and _is_empty_paragraph_el(children[start_idx]):
            body.remove(children[start_idx])
            children = list(body.iterchildren())

        # если уже стоит page break — не дублируем
        children = list(body.iterchildren())
        if start_idx < len(children):
            el = children[start_idx]
            if el.tag == qn("w:p"):
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                br = el.find(".//w:br", namespaces=ns)
                if br is not None and br.get(qn("w:type")) == "page":
                    continue

        _insert_page_break_at_body_index(doc, start_idx)


def _drop_unused_sections_by_tables(doc: Document, keep_sections: list[bool]) -> None:
    """
    В шаблоне блоки идут по таблицам:
      table0 -> смывы
      table1 -> воздух
      table2 -> мазки (персонал)

    Удаляем блок целиком (заголовки/абзацы + таблица),
    если по категории нет данных — чтобы не оставлять пустых "актов".

    keep_sections = [keep_swabs, keep_air, keep_smears]
    """
    body = doc._element.body  # pylint: disable=protected-access
    children = list(body.iterchildren())

    # preserve sectPr (чтобы не слетала ориентация)
    sectPr_copy = None
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr_copy = deepcopy(sectPr)
    else:
        for el in reversed(children):
            if el.tag == qn("w:p"):
                pPr = el.find(qn("w:pPr"))
                if pPr is not None:
                    sp = pPr.find(qn("w:sectPr"))
                    if sp is not None:
                        sectPr_copy = deepcopy(sp)
                        break

    # positions of each table element in body children
    tbl_els = [t._element for t in doc.tables]
    tbl_pos: list[int] = []
    for tbl_el in tbl_els:
        for idx, el in enumerate(children):
            if el is tbl_el:
                tbl_pos.append(idx)
                break

    n_tables = len(tbl_pos)
    n_sections = min(len(keep_sections), n_tables)

    # ranges for each section: start..end (inclusive), end includes the table itself
    ranges: list[tuple[int, int]] = []
    for k in range(n_sections):
        start = 0 if k == 0 else (tbl_pos[k - 1] + 1)
        end = tbl_pos[k]
        ranges.append((start, end))

    # delete from bottom to top so indexes don't shift
    for k in range(n_sections - 1, -1, -1):
        if keep_sections[k]:
            continue
        start, end = ranges[k]
        for el in children[start:end + 1]:
            if el.getparent() is body:
                body.remove(el)

    # restore sectPr if it disappeared
    if body.find(qn("w:sectPr")) is None and sectPr_copy is not None:
        body.append(sectPr_copy)


# -------------------------
# main entry
# -------------------------

def build_docx_report(
    template_path: str,
    out_path: str,
    dep: str,
    day: pd.Timestamp,
    df_swabs: pd.DataFrame,
    df_air: pd.DataFrame,
    df_smears: pd.DataFrame,
    *,
    auto_word_dir: Optional[str] = None,
    auto_pdf_dir: Optional[str] = None,
) -> Tuple[str, Optional[str], Optional[str]]:
    """
    Возвращает:
      (saved_docx_path, auto_word_docx_path_or_None, auto_pdf_path_or_None)

    Поведение:
    - создаём DOCX по шаблону и сохраняем в out_path
    - дополнительно копируем DOCX в AUTO_WORD_DIR (или auto_word_dir, если задан)
    - дополнительно сохраняем PDF в AUTO_PDF_DIR (или auto_pdf_dir, если задан)
    - если категории нет (пустой df) — её блок (акт+таблица) вырезается целиком
    - каждый оставшийся акт начинается с новой страницы (без "дыр" сверху)
    - ориентация всегда LANDSCAPE
    """
    date_str = _fmt_date_ddmmyyyy(day)

    doc = Document(template_path)
    _enforce_landscape(doc)

    # дата/отделение во всех блоках (смывы/воздух/персонал)
    _set_sampling_date_and_dep(doc, dep=dep, date_str=date_str)

    has_swabs = df_swabs is not None and not df_swabs.empty
    has_air = df_air is not None and not df_air.empty
    has_smears = df_smears is not None and not df_smears.empty

    if not (has_swabs or has_air or has_smears):
        raise ValueError("Нет данных для отчёта: смывы/воздух/персонал пустые.")

    # 1) убираем пустые блоки в шаблоне
    _drop_unused_sections_by_tables(doc, keep_sections=[has_swabs, has_air, has_smears])

    # 2) каждый оставшийся акт — с новой страницы (кроме первого), + чистим пустые абзацы
    _start_each_act_on_new_page(doc)

    # 3) заполняем таблицы по порядку оставшихся
    table_idx = 0
    if has_swabs and len(doc.tables) > table_idx:
        _fill_swabs_table(doc.tables[table_idx], df_swabs)
        table_idx += 1

    if has_air and len(doc.tables) > table_idx:
        _fill_air_table(doc.tables[table_idx], df_air)
        table_idx += 1

    if has_smears and len(doc.tables) > table_idx:
        _fill_smears_table(doc.tables[table_idx], df_smears)
        table_idx += 1

    _enforce_landscape(doc)

    # --- сохраняем основной DOCX ---
    saved_docx = os.path.normpath(os.path.abspath(str(out_path)))
    Path(saved_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(saved_docx)

    auto_word_docx = None
    auto_pdf = None

    # --- автосохранение Word (копия DOCX в сетевую папку) ---
    try:
        word_dir = auto_word_dir or AUTO_WORD_DIR
        Path(word_dir).mkdir(parents=True, exist_ok=True)
        auto_word_docx = os.path.normpath(str(Path(word_dir) / Path(saved_docx).name))
        shutil.copy2(saved_docx, auto_word_docx)
    except Exception as e:
        print(f"Не удалось сохранить Word-копию: {e}")
        auto_word_docx = None

    # --- автосохранение PDF (в сетевую папку) ---
    try:
        pdf_dir = auto_pdf_dir or AUTO_PDF_DIR
        Path(pdf_dir).mkdir(parents=True, exist_ok=True)
        auto_pdf = os.path.normpath(str(Path(pdf_dir) / (Path(saved_docx).stem + ".pdf")))
        save_pdf_via_word(saved_docx, auto_pdf)
    except Exception as e:
        print(f"Не удалось сохранить PDF: {e}")
        auto_pdf = None

    return saved_docx, auto_word_docx, auto_pdf
